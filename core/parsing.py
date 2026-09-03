"""Document parsing: turns raw files into (location_label, text) sections.

This is the local, free stand-in for a hosted OCR step. pymupdf handles
normal digital PDFs directly. For scanned/image-only pages and for
standalone photos (JPG/PNG/HEIC -- the "Handyfoto vom Vertrag" case), we
fall back to a local Tesseract OCR pass -- everything downstream (chunking,
embedding, search) stays unchanged either way.
"""
from __future__ import annotations

import re
from pathlib import Path

import fitz  # pymupdf

Section = tuple[str, str]  # (location_label, text)

# Below this many characters, a page is treated as "no usable text layer"
# (e.g. a scanned page) and gets OCR'd instead.
MIN_TEXT_LAYER_CHARS = 20
OCR_DPI = 300
OCR_LANGUAGES = "deu+eng"

_HEIF_REGISTERED = False


def _register_heif() -> None:
    global _HEIF_REGISTERED
    if not _HEIF_REGISTERED:
        try:
            import pillow_heif

            pillow_heif.register_heif_opener()
        except ImportError:
            pass
        _HEIF_REGISTERED = True


def _ocr_image(image) -> str:
    try:
        import pytesseract
    except ImportError:
        return ""
    try:
        return pytesseract.image_to_string(image, lang=OCR_LANGUAGES).strip()
    except pytesseract.TesseractNotFoundError:
        return ""


def _ocr_page(page: fitz.Page) -> str:
    try:
        from PIL import Image
    except ImportError:
        return ""

    zoom = OCR_DPI / 72
    pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
    return _ocr_image(image)


def parse_pdf(path: Path) -> list[Section]:
    doc = fitz.open(path)
    sections: list[Section] = []
    for i, page in enumerate(doc, start=1):
        text = page.get_text().strip()
        label = f"S. {i}"
        if len(text) < MIN_TEXT_LAYER_CHARS:
            ocr_text = _ocr_page(page)
            if len(ocr_text) > len(text):
                text, label = ocr_text, f"S. {i} (OCR)"
        if text:
            sections.append((label, text))
    return sections


def parse_image(path: Path) -> list[Section]:
    _register_heif()
    from PIL import Image

    with Image.open(path) as image:
        text = _ocr_image(image.convert("RGB"))
    return [(f"{path.name} (OCR)", text)] if text else []


def parse_docx(path: Path) -> list[Section]:
    import docx

    doc = docx.Document(str(path))
    sections: list[Section] = []
    heading, buffer = "Einleitung", []

    def flush():
        body = "\n\n".join(buffer).strip()
        if body:
            sections.append((heading, body))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading") or para.style.name == "Title":
            flush()
            heading, buffer = text, []
        else:
            buffer.append(text)
    flush()

    for i, table in enumerate(doc.tables, start=1):
        rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        table_text = "\n".join(r for r in rows if r.strip())
        if table_text:
            sections.append((f"Tabelle {i}", table_text))

    return sections


_HEADING_RE = re.compile(r"^(#{1,3}\s+.+)$", flags=re.MULTILINE)


def parse_markdown(path: Path) -> list[Section]:
    text = path.read_text(encoding="utf-8")
    parts = _HEADING_RE.split(text)
    sections: list[Section] = []
    if parts[0].strip():
        sections.append(("Einleitung", parts[0].strip()))
    for i in range(1, len(parts), 2):
        heading = parts[i].lstrip("#").strip()
        body = parts[i + 1].strip() if i + 1 < len(parts) else ""
        if body:
            sections.append((heading, body))
    return sections


def parse_txt(path: Path) -> list[Section]:
    text = path.read_text(encoding="utf-8").strip()
    return [("", text)] if text else []


PARSERS = {
    ".pdf": parse_pdf,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_txt,
    ".docx": parse_docx,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".png": parse_image,
    ".heic": parse_image,
    ".heif": parse_image,
}


def parse_document(path: Path) -> list[Section]:
    parser = PARSERS.get(path.suffix.lower())
    if parser is None:
        raise ValueError(
            f"Kein Parser für Dateityp '{path.suffix}' ({path.name}). "
            f"Unterstützt: {', '.join(PARSERS)}"
        )
    return parser(path)
